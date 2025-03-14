# Import Libraries
import os
import docx
import PyPDF2
import tempfile
import logging
from pathlib import Path
import google.generativeai as genai
from langchain_community.vectorstores import FAISS
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.chains import ConversationalRetrievalChain
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from typing import Optional, Union, BinaryIO, Dict, List, Tuple, Any

# Configure logging for FAISS
logging.getLogger("faiss").setLevel(logging.ERROR)  # Suppress FAISS warnings

# Define RAGProcessor Class
class RAGProcessor:
    """
    A processor for Retrieval-Augmented Generation (RAG) using Google Gemini models.
    """

    def __init__(self, api_key: Optional[str] = None) -> None:
        """
        Initialize the RAG processor with Google Gemini Flash 2.0.

        Args:
            api_key: Google API key. If None, it will try to get from environment variable.

        Raises:
            ValueError: If no API key is provided or found in environment variables.
        """
        self.api_key = api_key or os.environ.get("GOOGLE_API_KEY")
        if not self.api_key:
            raise ValueError(
                "Google API key is required. Set it as an environment variable or pass it directly."
            )

        # Configure the Google Generative AI
        genai.configure(api_key=self.api_key)

        # Initialize embeddings model
        self.embeddings = GoogleGenerativeAIEmbeddings(
            model="models/embedding-001", google_api_key=self.api_key
        )

        # Initialize LLM - using the latest Gemini model
        self.llm = ChatGoogleGenerativeAI(
            model="gemini-2.0-flash",  # Updated to the latest model
            google_api_key=self.api_key,
            temperature=0.2,
            top_p=0.95,
            top_k=40,
            max_output_tokens=2048,
        )

        # Text splitter for chunking documents
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,  # Increased overlap for better context
            separators=["\n\n", "\n", ". ", " ", ""],
        )

        self.vector_store = None
        self.retrieval_chain = None

    def extract_text_from_pdf(self, pdf_file: BinaryIO) -> str:
        """
        Extract text from a PDF file.

        Args:
            pdf_file: The uploaded PDF file object

        Returns:
            Extracted text from the PDF

        Raises:
            Exception: If there's an error processing the PDF
        """
        # Create a temporary file to store the PDF
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_file:
            temp_file_path = temp_file.name
            # Save the uploaded file to the temporary file
            pdf_file.seek(0)
            temp_file.write(pdf_file.read())

        try:
            # Extract text from the PDF
            text = ""
            with open(temp_file_path, "rb") as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + "\n\n"
            return text
        finally:
            # Clean up the temporary file
            Path(temp_file_path).unlink(missing_ok=True)

    def extract_text_from_docx(self, docx_file: BinaryIO) -> str:
        """
        Extract text from a Word document.

        Args:
            docx_file: The uploaded Word document file object

        Returns:
            Extracted text from the Word document

        Raises:
            Exception: If there's an error processing the Word document
        """
        # Create a temporary file to store the Word document
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_file:
            temp_file_path = temp_file.name
            # Save the uploaded file to the temporary file
            docx_file.seek(0)
            temp_file.write(docx_file.read())

        try:
            # Extract text from the Word document
            doc = docx.Document(temp_file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
                text += "\n"

            return text
        finally:
            # Clean up the temporary file
            Path(temp_file_path).unlink(missing_ok=True)

    def process_document(
        self,
        file: Optional[BinaryIO] = None,
        file_type: Optional[str] = None,
        text: Optional[str] = None,
    ) -> bool:
        """
        Process a document for RAG.

        Args:
            file: The uploaded file object (optional)
            file_type: The type of the file ('pdf' or 'docx') (optional)
            text: Text content to process (optional, used for OCR output)

        Returns:
            True if processing was successful

        Raises:
            ValueError: If neither file nor text is provided, or if the file type is not supported
            Exception: If there's an error processing the document
        """
        # Extract text from file if provided
        if file and file_type:
            match file_type.lower():
                case "pdf":
                    document_text = self.extract_text_from_pdf(file)
                case "docx":
                    document_text = self.extract_text_from_docx(file)
                case _:
                    raise ValueError(f"Unsupported file type for RAG: {file_type}")
        elif text:
            document_text = text
        else:
            raise ValueError("Either file or text must be provided")

        # Split text into chunks
        chunks = self.text_splitter.split_text(document_text)

        # Create vector store
        self.vector_store = FAISS.from_texts(chunks, self.embeddings)

        # Create retrieval chain with improved configuration
        self.retrieval_chain = ConversationalRetrievalChain.from_llm(
            llm=self.llm,
            retriever=self.vector_store.as_retriever(
                search_type="mmr",  # Use Maximum Marginal Relevance for better diversity
                search_kwargs={
                    "k": 5,
                    "fetch_k": 10,
                },  # Retrieve more documents for better context
            ),
            return_source_documents=True,
            verbose=True,
        )

        return True

    def query(
        self, query: str, chat_history: Optional[List[Tuple[str, str]]] = None
    ) -> Dict[str, Any]:
        """
        Query the RAG system.

        Args:
            query: The user's query
            chat_history: List of previous interactions as (human_message, ai_message) tuples

        Returns:
            Response from the RAG system containing answer and source documents

        Raises:
            ValueError: If no document has been processed
            Exception: If there's an error querying the system
        """
        if not self.retrieval_chain:
            raise ValueError(
                "No document has been processed. Please process a document first."
            )

        chat_history = chat_history or []

        # Query the retrieval chain
        response = self.retrieval_chain(
            {"question": query, "chat_history": chat_history}
        )

        return {
            "answer": response["answer"],
            "source_documents": response.get("source_documents", []),
        }
